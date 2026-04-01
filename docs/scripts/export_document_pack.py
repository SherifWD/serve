#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
HTML_DIR = ROOT / "docs" / "html"
EXPORT_DOCX_DIR = ROOT / "docs" / "exports" / "docx"
EXPORT_PDF_DIR = ROOT / "docs" / "exports" / "pdf"

HEADING_TAGS = {"h1", "h2", "h3"}
CONTAINER_TAGS = {"body", "div", "section", "article", "main"}
SKIP_TAGS = {"style", "script", "link", "meta", "head"}
MAJOR_BLOCK_TAGS = {"h1", "h2", "h3", "p", "ul", "ol", "table", "img", "div", "section", "article"}


def normalize_space(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def inline_text(node: Tag | NavigableString) -> str:
    if isinstance(node, NavigableString):
        return str(node)

    if node.name == "br":
        return "\n"

    if node.name == "a":
        label = normalize_space(" ".join(inline_text(child) for child in node.children))
        href = node.get("href", "").strip()
        if href and label:
            return f"{label} ({href})"
        return label or href

    if node.name == "code":
        return f"`{normalize_space(node.get_text(' ', strip=True))}`"

    parts = [inline_text(child) for child in node.children]
    return normalize_space(" ".join(part for part in parts if part))


def has_major_children(tag: Tag) -> bool:
    return any(isinstance(child, Tag) and child.name in MAJOR_BLOCK_TAGS for child in tag.children)


def resolve_image_path(tag: Tag, base_dir: Path) -> Path | None:
    src = (tag.get("src") or "").strip()
    if not src or src.startswith(("http://", "https://", "data:")):
        return None
    path = (base_dir / src).resolve()
    return path if path.exists() else None


def image_dimensions(path: Path, max_width_in: float, max_height_in: float) -> tuple[float, float]:
    with PILImage.open(path) as image:
        width_px, height_px = image.size

    dpi = 96
    width_in = width_px / dpi
    height_in = height_px / dpi
    scale = min(max_width_in / width_in, max_height_in / height_in, 1.0)
    return width_in * scale, height_in * scale


def table_matrix(tag: Tag) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in tag.find_all("tr", recursive=True):
        cells = row.find_all(["th", "td"], recursive=False) or row.find_all(["th", "td"], recursive=True)
        parsed = [inline_text(cell) for cell in cells]
        if any(parsed):
            rows.append(parsed)
    return rows


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", parent=styles["Title"], fontSize=20, leading=24, spaceAfter=16))
    styles.add(ParagraphStyle(name="DocH1", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=12))
    styles.add(ParagraphStyle(name="DocH2", parent=styles["Heading2"], fontSize=15, leading=19, spaceBefore=10, spaceAfter=8))
    styles.add(ParagraphStyle(name="DocH3", parent=styles["Heading3"], fontSize=12, leading=16, spaceBefore=8, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontSize=10.5, leading=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="Caption", parent=styles["BodyText"], fontSize=8.5, textColor=colors.HexColor("#555555"), leading=11, spaceAfter=8))
    styles.add(ParagraphStyle(name="Note", parent=styles["BodyText"], fontSize=9.5, textColor=colors.HexColor("#444444"), leading=13, spaceAfter=8))
    return styles


def append_pdf_flowables(tag: Tag, base_dir: Path, flowables: list, styles) -> None:
    if tag.name in SKIP_TAGS:
        return

    classes = set(tag.get("class", []))

    if tag.name == "img":
        path = resolve_image_path(tag, base_dir)
        if not path:
            return
        width_in, height_in = image_dimensions(path, max_width_in=6.8, max_height_in=7.0)
        flowables.append(RLImage(str(path), width=width_in * inch, height=height_in * inch))
        flowables.append(Spacer(1, 0.12 * inch))
        return

    if tag.name in HEADING_TAGS:
        text = inline_text(tag)
        if not text:
            return
        style_name = {"h1": "DocTitle", "h2": "DocH1", "h3": "DocH2"}[tag.name]
        flowables.append(Paragraph(escape(text), styles[style_name]))
        return

    if tag.name == "p":
        text = inline_text(tag)
        if not text:
            return
        flowables.append(Paragraph(escape(text), styles["Body"]))
        return

    if tag.name in {"ul", "ol"}:
        items = []
        bullet_type = "bullet" if tag.name == "ul" else "1"
        for item in tag.find_all("li", recursive=False):
            item_text = inline_text(item)
            if item_text:
                items.append(ListItem(Paragraph(escape(item_text), styles["Body"])))
        if items:
            flowables.append(ListFlowable(items, bulletType=bullet_type, leftIndent=18))
            flowables.append(Spacer(1, 0.08 * inch))
        return

    if tag.name == "table":
        rows = table_matrix(tag)
        if not rows:
            return
        max_cols = max(len(row) for row in rows)
        padded_rows = [row + [""] * (max_cols - len(row)) for row in rows]
        col_width = 6.8 * inch / max_cols
        table = Table(padded_rows, colWidths=[col_width] * max_cols, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2d39c")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#b8b8b8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("LEADING", (0, 0), (-1, -1), 11),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fcf7ef")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        flowables.append(table)
        flowables.append(Spacer(1, 0.14 * inch))
        return

    if tag.name in CONTAINER_TAGS:
        if has_major_children(tag):
            for child in tag.children:
                if isinstance(child, Tag):
                    append_pdf_flowables(child, base_dir, flowables, styles)
            return

        text = inline_text(tag)
        if text:
            style_name = "Caption" if "caption" in classes else "Note" if {"note", "pill-row"} & classes else "Body"
            flowables.append(Paragraph(escape(text), styles[style_name]))
        return

    text = inline_text(tag)
    if text:
        flowables.append(Paragraph(escape(text), styles["Body"]))


def build_pdf(html_path: Path, output_path: Path) -> None:
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "lxml")
    flowables = []
    styles = pdf_styles()

    body = soup.body or soup
    for child in body.children:
        if isinstance(child, Tag):
            append_pdf_flowables(child, html_path.parent, flowables, styles)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title=(soup.title.string if soup.title else html_path.stem),
    )
    document.build(flowables)


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)


def add_docx_paragraph(document: Document, text: str, *, style: str | None = None, bold: bool = False, italic: bool = False) -> None:
    if not text:
        return
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def append_docx_content(tag: Tag, base_dir: Path, document: Document) -> None:
    if tag.name in SKIP_TAGS:
        return

    classes = set(tag.get("class", []))

    if tag.name in HEADING_TAGS:
        text = inline_text(tag)
        if not text:
            return
        level = {"h1": 0, "h2": 1, "h3": 2}[tag.name]
        document.add_heading(text, level=level)
        return

    if tag.name == "p":
        add_docx_paragraph(document, inline_text(tag))
        return

    if tag.name in {"ul", "ol"}:
        style = "List Bullet" if tag.name == "ul" else "List Number"
        for item in tag.find_all("li", recursive=False):
            item_text = inline_text(item)
            if item_text:
                document.add_paragraph(item_text, style=style)
        return

    if tag.name == "img":
        path = resolve_image_path(tag, base_dir)
        if not path:
            return
        width_in, _ = image_dimensions(path, max_width_in=6.0, max_height_in=8.0)
        document.add_picture(str(path), width=Inches(width_in))
        image_paragraph = document.paragraphs[-1]
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return

    if tag.name == "table":
        rows = table_matrix(tag)
        if not rows:
            return
        max_cols = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for column_index in range(max_cols):
                value = row[column_index] if column_index < len(row) else ""
                cell = table.rows[row_index].cells[column_index]
                cell.text = value
                if row_index == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        document.add_paragraph("")
        return

    if tag.name in CONTAINER_TAGS:
        if has_major_children(tag):
            for child in tag.children:
                if isinstance(child, Tag):
                    append_docx_content(child, base_dir, document)
            return

        text = inline_text(tag)
        if text:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            if "caption" in classes:
                run.italic = True
            if "note" in classes:
                run.bold = True
        return

    text = inline_text(tag)
    if text:
        add_docx_paragraph(document, text)


def build_docx(html_path: Path, output_path: Path) -> None:
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "lxml")
    document = Document()
    configure_docx_styles(document)

    body = soup.body or soup
    for child in body.children:
        if isinstance(child, Tag):
            append_docx_content(child, html_path.parent, document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def html_sources(paths: Iterable[Path]) -> list[Path]:
    return [path for path in paths if path.suffix == ".html" and path.name != "style.css"]


def main() -> int:
    parser = argparse.ArgumentParser(description="Export HTML documentation to PDF and DOCX.")
    parser.add_argument("html_files", nargs="*", help="Relative paths under docs/html. Defaults to all HTML docs.")
    args = parser.parse_args()

    if args.html_files:
        sources = [HTML_DIR / path for path in args.html_files]
    else:
        sources = sorted(html_sources(HTML_DIR.iterdir()))

    EXPORT_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    EXPORT_PDF_DIR.mkdir(parents=True, exist_ok=True)

    for html_path in sources:
        if not html_path.exists():
            raise FileNotFoundError(f"Missing HTML document: {html_path}")

        build_docx(html_path, EXPORT_DOCX_DIR / f"{html_path.stem}.docx")
        build_pdf(html_path, EXPORT_PDF_DIR / f"{html_path.stem}.pdf")
        print(f"Exported {html_path.name}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
